import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re

def convert():
    scripts_dir = os.path.dirname(os.path.abspath(__file__))
    project_dir = os.path.dirname(scripts_dir)
    docs_dir = os.path.join(project_dir, 'docs')
    md_path = os.path.join(docs_dir, 'Инструкция_Повторное_посещение.md')
    docx_path = os.path.join(docs_dir, 'Инструкция_Повторное_посещение.docx')
    
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return
        
    doc = docx.Document()
    
    # Configure styles
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line_str = line.strip()
        if not line_str:
            continue
            
        # Heading 1
        if line_str.startswith('# '):
            text = line_str[2:]
            doc.add_heading(text, level=1)
            
        # Heading 3
        elif line_str.startswith('### '):
            text = line_str[4:]
            doc.add_heading(text, level=3)
            
        # Bullet list
        elif line_str.startswith('* '):
            text = line_str[2:]
            p = doc.add_paragraph(style='List Bullet')
            
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
                    
        # Sub-bullets (nested lists, starting with '  *' or '  -')
        elif line_str.startswith('- ') or line.startswith('  * ') or line.startswith('  - '):
            text = line_str[2:]
            p = doc.add_paragraph(style='List Bullet 2')
            
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
                    
        # Separator line
        elif line_str == '---':
            p = doc.add_paragraph()
            p.add_run('____________________________________________________').font.color.rgb = docx.shared.RGBColor(200, 200, 200)
            
        # Image embed
        elif line_str.startswith('![') and ']' in line_str and '(' in line_str:
            match = re.search(r'!\[(.*?)\]\((.*?)\)', line_str)
            if match:
                caption = match.group(1)
                img_path = match.group(2)
                
                # Resolve path
                abs_img_path = img_path
                if not os.path.isabs(img_path):
                    abs_img_path = os.path.join(docs_dir, img_path)
                    
                # Skip video files and animations in the word document
                if img_path.endswith('.webm') or img_path.endswith('.mp4') or (img_path.endswith('.webp') and 'demo' in img_path):
                    p = doc.add_paragraph()
                    run = p.add_run(f'[Ссылка на видеозапись демонстрации: {os.path.basename(img_path)}]')
                    run.italic = True
                    run.font.color.rgb = docx.shared.RGBColor(81, 35, 212)
                    continue

                if os.path.exists(abs_img_path):
                    try:
                        p_img = doc.add_paragraph()
                        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_img.add_run().add_picture(abs_img_path, width=Inches(5.5))
                        
                        p_cap = doc.add_paragraph()
                        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run_cap = p_cap.add_run(f'Рис. {caption}')
                        run_cap.italic = True
                        run_cap.font.size = Pt(9.5)
                        run_cap.font.color.rgb = docx.shared.RGBColor(120, 120, 120)
                    except Exception as e:
                        print(f"Error adding image {abs_img_path}: {e}")
                else:
                    print(f"Image path not found: {abs_img_path}")
                    
        # Normal text
        else:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line_str)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
                    
    doc.save(docx_path)
    print("DOCX compiled successfully at:", docx_path)

if __name__ == '__main__':
    convert()
